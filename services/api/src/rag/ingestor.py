"""
Ingestor de documentos para RAG
"""

import os
import io
from typing import List, Optional, Dict, Any
from pathlib import Path
import httpx
import structlog

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from .vectorstore import RAGVectorStore
from .config import get_rag_config

logger = structlog.get_logger()


class DocumentIngestor:
    """Ingestor de documentos con soporte para múltiples formatos"""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".html", ".xlsx", ".xls", ".csv"}
    
    def __init__(
        self,
        collection: str = "default",
        chunk_size: int = None,
        chunk_overlap: int = None,
        embedding_model: str = None,
        embedding_base_url: str = None
    ):
        config = get_rag_config()
        
        self.collection = collection
        self.chunk_size = chunk_size or config.chunk_size
        self.chunk_overlap = chunk_overlap or config.chunk_overlap
        
        self.vectorstore = RAGVectorStore(
            collection=collection,
            embedding_model=embedding_model,
            embedding_base_url=embedding_base_url
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    async def ingest_file(
        self,
        file_path: str,
        document_id: str = None,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Ingestar un archivo local"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Formato no soportado: {extension}")
        
        # Extraer texto según el formato
        text = await self._extract_text(path, extension)
        
        if not text.strip():
            raise ValueError("El documento está vacío o no se pudo extraer texto")
        
        # Preparar metadatos
        doc_metadata = {
            "source": path.name,
            "file_path": str(path),
            "file_type": extension,
            "file_size": path.stat().st_size,
            **(metadata or {})
        }
        
        # Dividir en chunks
        chunks = self.text_splitter.split_text(text)
        
        # Preparar metadatos para cada chunk
        chunk_metadatas = [
            {**doc_metadata, "chunk_index": i, "total_chunks": len(chunks)}
            for i in range(len(chunks))
        ]
        
        # Añadir al vector store
        doc_id = document_id or path.stem
        chunk_ids = await self.vectorstore.add_documents(
            texts=chunks,
            metadatas=chunk_metadatas,
            document_id=doc_id
        )
        
        return {
            "document_id": doc_id,
            "file_name": path.name,
            "chunks_created": len(chunks),
            "chunk_ids": chunk_ids
        }
    
    async def ingest_from_url(
        self,
        url: str,
        document_id: str = None,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Ingestar documento desde URL (Strapi u otra fuente)"""
        logger.info(f"Descargando documento desde {url}")
        
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.get(url)
            response.raise_for_status()
            content = response.content
        
        # Detectar tipo de archivo
        content_type = response.headers.get("content-type", "")
        file_name = url.split("/")[-1].split("?")[0]
        
        extension = self._detect_extension(content_type, file_name)
        
        # Extraer texto
        text = await self._extract_text_from_bytes(content, extension)
        
        if not text.strip():
            raise ValueError("El documento está vacío o no se pudo extraer texto")
        
        # Preparar metadatos
        doc_metadata = {
            "source": file_name,
            "source_url": url,
            "file_type": extension,
            "file_size": len(content),
            **(metadata or {})
        }
        
        # Dividir en chunks
        chunks = self.text_splitter.split_text(text)
        
        chunk_metadatas = [
            {**doc_metadata, "chunk_index": i, "total_chunks": len(chunks)}
            for i in range(len(chunks))
        ]
        
        # Añadir al vector store
        doc_id = document_id or file_name.rsplit(".", 1)[0]
        chunk_ids = await self.vectorstore.add_documents(
            texts=chunks,
            metadatas=chunk_metadatas,
            document_id=doc_id
        )
        
        return {
            "document_id": doc_id,
            "file_name": file_name,
            "chunks_created": len(chunks),
            "chunk_ids": chunk_ids
        }
    
    async def ingest_text(
        self,
        text: str,
        document_id: str,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Ingestar texto directamente"""
        if not text.strip():
            raise ValueError("El texto está vacío")
        
        doc_metadata = {
            "source": "direct_text",
            "file_type": ".txt",
            **(metadata or {})
        }
        
        chunks = self.text_splitter.split_text(text)
        
        chunk_metadatas = [
            {**doc_metadata, "chunk_index": i, "total_chunks": len(chunks)}
            for i in range(len(chunks))
        ]
        
        chunk_ids = await self.vectorstore.add_documents(
            texts=chunks,
            metadatas=chunk_metadatas,
            document_id=document_id
        )
        
        return {
            "document_id": document_id,
            "chunks_created": len(chunks),
            "chunk_ids": chunk_ids
        }
    
    async def delete_document(self, document_id: str) -> int:
        """Eliminar un documento del índice"""
        return await self.vectorstore.delete_by_document(document_id)
    
    async def _extract_text(self, path: Path, extension: str) -> str:
        """Extraer texto de un archivo"""
        if extension == ".txt" or extension == ".md":
            return path.read_text(encoding="utf-8")
        
        elif extension == ".pdf":
            return await self._extract_pdf(path)
        
        elif extension == ".docx":
            return await self._extract_docx(path)
        
        elif extension == ".html":
            return await self._extract_html(path.read_text(encoding="utf-8"))
        
        elif extension == ".csv":
            return await self._extract_csv(path)
        
        elif extension in (".xlsx", ".xls"):
            return await self._extract_excel(path)
        
        else:
            raise ValueError(f"Extractor no implementado para {extension}")
    
    async def _extract_text_from_bytes(self, content: bytes, extension: str) -> str:
        """Extraer texto de bytes"""
        if extension in [".txt", ".md"]:
            return content.decode("utf-8")
        
        elif extension == ".pdf":
            return await self._extract_pdf_bytes(content)
        
        elif extension == ".docx":
            return await self._extract_docx_bytes(content)
        
        elif extension == ".html":
            return await self._extract_html(content.decode("utf-8"))
        
        elif extension == ".csv":
            return await self._extract_csv_bytes(content)
        
        elif extension in (".xlsx", ".xls"):
            return await self._extract_excel_bytes(content)
        
        else:
            raise ValueError(f"Extractor no implementado para {extension}")
    
    async def _extract_pdf(self, path: Path) -> str:
        """Extraer texto de PDF"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pypdf no instalado. Ejecuta: pip install pypdf")
    
    async def _extract_pdf_bytes(self, content: bytes) -> str:
        """Extraer texto de PDF desde bytes"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pypdf no instalado")
    
    async def _extract_docx(self, path: Path) -> str:
        """Extraer texto de DOCX"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            raise ImportError("python-docx no instalado")
    
    async def _extract_docx_bytes(self, content: bytes) -> str:
        """Extraer texto de DOCX desde bytes"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            raise ImportError("python-docx no instalado")
    
    async def _extract_html(self, html: str) -> str:
        """Extraer texto de HTML"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            # Eliminar scripts y estilos
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            # Fallback simple sin BeautifulSoup
            import re
            text = re.sub(r'<[^>]+>', ' ', html)
            return re.sub(r'\s+', ' ', text).strip()
    
    async def _extract_csv(self, path: Path) -> str:
        """Extraer texto de CSV preservando estructura tabular"""
        import csv
        rows = []
        with open(path, newline='', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)

    async def _extract_csv_bytes(self, content: bytes) -> str:
        """Extraer texto de CSV desde bytes"""
        import csv
        text = content.decode("utf-8-sig")
        rows = []
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            rows.append(" | ".join(row))
        return "\n".join(rows)

    async def _extract_excel(self, path: Path) -> str:
        """Extraer texto de Excel preservando estructura tabular"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path), read_only=True, data_only=True)
            sheets_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(" | ".join(cells))
                if rows:
                    sheets_text.append(f"## Hoja: {sheet_name}\n\n" + "\n".join(rows))
            wb.close()
            return "\n\n".join(sheets_text)
        except ImportError:
            raise ImportError("openpyxl no instalado. Ejecuta: pip install openpyxl")

    async def _extract_excel_bytes(self, content: bytes) -> str:
        """Extraer texto de Excel desde bytes"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sheets_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(" | ".join(cells))
                if rows:
                    sheets_text.append(f"## Hoja: {sheet_name}\n\n" + "\n".join(rows))
            wb.close()
            return "\n\n".join(sheets_text)
        except ImportError:
            raise ImportError("openpyxl no instalado")

    def _detect_extension(self, content_type: str, file_name: str) -> str:
        """Detectar extensión del archivo"""
        # Por content-type
        type_map = {
            "application/pdf": ".pdf",
            "text/plain": ".txt",
            "text/markdown": ".md",
            "text/html": ".html",
            "text/csv": ".csv",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
            "application/vnd.ms-excel": ".xls"
        }
        
        for mime, ext in type_map.items():
            if mime in content_type:
                return ext
        
        # Por nombre de archivo
        if "." in file_name:
            return "." + file_name.rsplit(".", 1)[-1].lower()
        
        return ".txt"  # Default
